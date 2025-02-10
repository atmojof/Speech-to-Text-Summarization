# Library
import streamlit as st
import magic
import PyPDF2
from tqdm import tqdm
from docx import Document
from transformers import pipeline
import time
#from google.colab import files
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from fpdf import FPDF
from io import BytesIO

from transformers.models import pegasus 


# Streamlit Output
st.title("Long Text Summarizer")
st.header("with Bart Large CNN (facebook) and Pegasus CNN Dailymail (google)")
Jenis_Model = st.multiselect(
    "Select Model : ",
    ["Bart Large CNN (Extractive)", "Pegasus CNN Dailymail (Abstractive)", "Both"],
max_selections=1)

uploaded_file = st.file_uploader("Choose a file to summarize")

if uploaded_file is not None:
  mime_type = magic.Magic(mime=True).from_buffer(uploaded_file.read(1024))  # Membaca sebagian file untuk deteksi tipe
  st.write(f"Tipe MIME berdasarkan konten: {mime_type}")

# Input 
# PDF
def open_PDF (uploaded_file) : # Open the PDF file
  # Menginisialisasi string kosong
  text = ""
  pdf_reader = PyPDF2.PdfReader(uploaded_file)
  for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text += page.extract_text()
  return text
# docx 
def open_docx (uploaded_file) : # Mebaca file docx dan mengisi variable text
  doc = Document(uploaded_file)
    
  text = ""
  for para in doc.paragraphs:
    text += para.text + "\n" 
  return text
# txt
def open_txt (uploaded_file) : # Mebaca file docx dan mengisi variable text    
  text = uploaded_file.getvalue().decode("utf-8")
  return text
# Input Classifier
def input_classifier (mime_type,uploaded_file) : # Menentukan jenis file berdasarkan path yang diberikan
  Jenis_File = mime_type[-3:]
  # Memanggil function sesuai dengan path file
  if Jenis_File == "pdf" :
    text = open_PDF (uploaded_file)
  elif Jenis_File == "txt" :
    text = open_txt (uploaded_file)
  else :
    text = open_docx (uploaded_file)
  return text

# Bart Model
def bart_summarizer (text) :
  pipe = pipeline('summarization', model='facebook/bart-large-cnn')
  inputs = pipe.tokenizer(text, return_tensors="pt", truncation=False)
  summaries = []

  if len(inputs['input_ids'][0]) > 1024:  # Adjust the limit based on model's max tokens
      # Split the text into smaller chunks
      chunk_size = 1024  # Define your chunk size
      chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

      for chunk in chunks:
          pipe_out = pipe(chunk,max_length=130, min_length=30, do_sample=False)
          print (pipe_out[0]['summary_text'])
          summaries.append(pipe_out[0]['summary_text'])
  else:
      # If length is within limit, proceed with summarization
      pipe_out = pipe(text)
      summaries.append(pipe_out[0]['summary_text'])
  paragraph = "\n".join(summaries)

  return paragraph

# Pegasus Model
def pegasus_summarizer (text) :
  pipe = pipeline('summarization', model ='google/pegasus-cnn_dailymail')
  inputs = pipe.tokenizer(text, return_tensors="pt", truncation=False)
  summaries = []
  if len(inputs['input_ids'][0]) > 1024:  # Adjust the limit based on model's max tokens
      # Split the text into smaller chunks
      chunk_size = 1024  # Define your chunk size
      chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

      for chunk in chunks:
          pipe_out = pipe(chunk,max_length=130, min_length=30, do_sample=False)
          summaries.append(pipe_out[0]['summary_text'])
  else:
      # If length is within limit, proceed with summarization
      pipe_out = pipe(text)
      summaries.append(pipe_out[0]['summary_text'])
  paragraph = "\n".join(summaries)

  return paragraph

# Output 
# Docx
def Output_DOCX (Text) :
  doc = Document()

  # Menambahkan teks ke dalam file DOCX
  doc.add_paragraph(Text)
  docx_stream = BytesIO()
  doc.save(docx_stream)
  docx_stream.seek(0)
  return (docx_stream)

# PDF
def Output_PDF (Text) :
  # Membuat objek PDF
  pdf = FPDF()
  pdf.set_auto_page_break(auto=True, margin=15)
  pdf.add_page()
  pdf.set_font("Arial", size=12)

  # Menambahkan teks ke dalam file PDF
  pdf.multi_cell(0, 10, Text)

  # Menyimpan file PDF ke dalam byte stream
  pdf_output = pdf.output(dest='S').encode('latin1')
  return (pdf_output)

st.download_button(
    label="Download PDF",
    data= pdf_output,  # Data dalam format byte
    file_name="Summary.pdf",  # Nama file PDF yang diunduh
    mime="application/pdf"  # Tipe MIME untuk PDF
)

# Main 
text = input_classifier (mime_type,uploaded_file)
word_count = len(text.split())
st.write (f"Word Count in File :  {word_count}")

if Jenis_Model == "Bart Large CNN (Extractive)" :
  Summary = bart_summarizer (text)
  word_count_bart = len(Summary.split())
  st.write (f"Word Count After Summarization (Bart Large CNN Model):  {word_count_bart}")
elif Jenis_Model == "Pegasus CNN Dailymail (Abstractive)" :
  Summary = pegasus_summarizer (text)   
  word_count_pegasus = len(Summary.split())
  st.write (f"Word Count After Summarization (Pegasus CNN Dailymail Model):  {word_count_pegasus}")
else :
  Summary = "Bart model Summarization :" +"/n"
  Bart = bart_summarizer (text)
  word_count_bart = len(Bart.split())
  st.write (f"Word Count After Summarization (Bart Large CNN Model):  {word_count_bart}")
  Summary = Summary + Bart
  Pegasus = pegasus_summarizer (text)   
  word_count_pegasus = len(Pegasus.split())
  st.write (f"Word Count After Summarization (Pegasus CNN Dailymail Model):  {word_count_pegasus}")
  Summary = Summary + Pegasus

Summary_Docx = Output_DOCX (Summary)
Summary_PDF = Output_PDF (Summary)

st.download_button(
    label="Download Docx File",
    data= Summary_Docx,  # Data DOCX
    file_name="Summary.docx",  # Nama file yang akan diunduh
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"  # Tipe MIME untuk file DOCX
)

st.download_button(
    label="Download PDF File",
    data= Summary_PDF,  # Data dalam format byte
    file_name="Summary.pdf",  # Nama file PDF yang diunduh
    mime="application/pdf"  # Tipe MIME untuk PDF
)

st.download_button(
    label="Download Txt File",
    data=Summary,  # Isi file
    file_name="Summary.txt",  # Nama file yang akan diunduh
    mime="text/plain"  # Tipe MIME untuk file teks
)